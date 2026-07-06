"""Regression tests for the three hot paths:

1. DOCX template import  (/templates/import)
2. Plan assignment       (/assign)
3. Task inline editing   (add task, status/notes/due-date HTMX flows)

These pin down the *current live behavior* in app.py so the blueprint
consolidation (and future changes) can be verified against them.
"""

import io
import os
import tempfile
from datetime import date

import docx
import pytest

from app import app
from Onboarding.extensions import db
from Onboarding.models import (
    OnboardingPlan,
    OnboardingTemplate,
    RoleEnum,
    StatusEnum,
    Task,
    TemplateSection,
    TemplateStatusEnum,
    TemplateTask,
    User,
    Week,
)


# ---------------------------------------------------------------------------
# Fixtures / helpers
# ---------------------------------------------------------------------------
@pytest.fixture(autouse=True)
def fresh_db():
    """Point the app at a throwaway SQLite file for each test."""
    fd, path = tempfile.mkstemp(suffix=".sqlite3")
    os.close(fd)
    with app.app_context():
        app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{path}"
        db.engine.dispose()
        db.drop_all()
        db.create_all()
    yield
    try:
        os.unlink(path)
    except OSError:
        pass


def make_user(email, name, role, manager_id=None):
    u = User(email=email, full_name=name, role=role, manager_id=manager_id)
    db.session.add(u)
    db.session.flush()
    return u


def as_user(email):
    """Identity header understood by current_user()."""
    return {"X-User-Email": email}


def make_docx(tables):
    """Build an in-memory .docx.

    ``tables`` is a list of dicts:
      {"title": "Week 1 - Orientation" | None,
       "rows": [("Task title", "Outcome"), ...]}
    Each table gets a Training/Outcomes header row, matching the format
    the importer parses.
    """
    d = docx.Document()
    for spec in tables:
        title = spec.get("title")
        rows = spec["rows"]
        n_rows = len(rows) + (2 if title else 1)
        table = d.add_table(rows=n_rows, cols=2)
        r = 0
        if title:
            table.rows[r].cells[0].text = title
            r += 1
        table.rows[r].cells[0].text = "Training"
        table.rows[r].cells[1].text = "Outcomes"
        r += 1
        for t_title, t_desc in rows:
            table.rows[r].cells[0].text = t_title
            table.rows[r].cells[1].text = t_desc
            r += 1
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


def make_published_template(created_by_id, n_sections=2, tasks_per_section=2):
    tpl = OnboardingTemplate(
        name="Warranty Rep Template",
        description="test",
        status=TemplateStatusEnum.PUBLISHED.value,
        created_by_id=created_by_id,
    )
    db.session.add(tpl)
    db.session.flush()
    for s in range(1, n_sections + 1):
        sec = TemplateSection(
            template_id=tpl.id,
            title=f"Week {s}",
            offset_days=(s - 1) * 7,
            order_index=s,
        )
        db.session.add(sec)
        db.session.flush()
        for t in range(1, tasks_per_section + 1):
            db.session.add(
                TemplateTask(
                    section_id=sec.id,
                    title=f"S{s} Task {t}",
                    description="desc",
                    order_index=t,
                    section_day=1,
                )
            )
    db.session.commit()
    return tpl


# ---------------------------------------------------------------------------
# 1. DOCX template import
# ---------------------------------------------------------------------------
class TestDocxImport:
    def test_builder_import_creates_sections_and_tasks(self):
        with app.app_context():
            make_user("builder@example.com", "Bea Builder", RoleEnum.BUILDER.value)
            db.session.commit()

        f = make_docx(
            [
                {
                    "title": "Week 1 - Orientation",
                    "rows": [
                        ("Meet the team", "Know your teammates"),
                        ("Laptop setup", "Working dev machine"),
                    ],
                },
                {
                    "title": "Week 3 - Systems",
                    "rows": [("CRM training", "Can log a case")],
                },
            ]
        )

        client = app.test_client()
        resp = client.post(
            "/templates/import",
            headers=as_user("builder@example.com"),
            data={"file": (f, "rep_plan.docx")},
            content_type="multipart/form-data",
        )
        assert resp.status_code == 302
        assert "/preview" in resp.headers["Location"]

        with app.app_context():
            tpl = OnboardingTemplate.query.one()
            assert tpl.name == "Imported: rep_plan"
            assert tpl.status == TemplateStatusEnum.DRAFT.value

            sections = (
                TemplateSection.query.filter_by(template_id=tpl.id)
                .order_by(TemplateSection.order_index)
                .all()
            )
            assert [s.title for s in sections] == [
                "Week 1 - Orientation",
                "Week 3 - Systems",
            ]
            # "Week N" in the title drives the offset: (N-1) * 7
            assert [s.offset_days for s in sections] == [0, 14]

            tasks_by_section = [
                TemplateTask.query.filter_by(section_id=s.id)
                .order_by(TemplateTask.order_index)
                .all()
                for s in sections
            ]
            assert [t.title for t in tasks_by_section[0]] == [
                "Meet the team",
                "Laptop setup",
            ]
            assert tasks_by_section[0][0].description == "Know your teammates"
            assert [t.title for t in tasks_by_section[1]] == ["CRM training"]

    def test_table_without_training_outcomes_header_is_skipped(self):
        with app.app_context():
            make_user("builder@example.com", "Bea Builder", RoleEnum.BUILDER.value)
            db.session.commit()

        # Header row missing "Outcomes" -> importer should skip the table
        d = docx.Document()
        table = d.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Stuff"
        table.rows[0].cells[1].text = "Things"
        table.rows[1].cells[0].text = "Not a task"
        buf = io.BytesIO()
        d.save(buf)
        buf.seek(0)

        client = app.test_client()
        resp = client.post(
            "/templates/import",
            headers=as_user("builder@example.com"),
            data={"file": (buf, "empty.docx")},
            content_type="multipart/form-data",
        )
        assert resp.status_code == 302  # template shell still created

        with app.app_context():
            tpl = OnboardingTemplate.query.one()
            assert TemplateSection.query.filter_by(template_id=tpl.id).count() == 0

    def test_import_rejects_non_builder(self):
        with app.app_context():
            make_user("user@example.com", "Uma User", RoleEnum.USER.value)
            db.session.commit()

        client = app.test_client()
        resp = client.post(
            "/templates/import",
            headers=as_user("user@example.com"),
            data={"file": (make_docx([]), "x.docx")},
            content_type="multipart/form-data",
        )
        assert resp.status_code == 403

    def test_import_rejects_wrong_extension(self):
        with app.app_context():
            make_user("builder@example.com", "Bea Builder", RoleEnum.BUILDER.value)
            db.session.commit()

        client = app.test_client()
        resp = client.post(
            "/templates/import",
            headers=as_user("builder@example.com"),
            data={"file": (io.BytesIO(b"nope"), "plan.txt")},
            content_type="multipart/form-data",
        )
        assert resp.status_code == 400


# ---------------------------------------------------------------------------
# 2. Plan assignment
# ---------------------------------------------------------------------------
class TestAssignPlan:
    def _setup_admin_tpl_employee(self):
        admin = make_user("admin@example.com", "Ada Admin", RoleEnum.ADMIN.value)
        emp = make_user("new@example.com", "Nia Newhire", RoleEnum.USER.value)
        tpl = make_published_template(admin.id, n_sections=2, tasks_per_section=2)
        db.session.commit()
        return admin, emp, tpl

    def test_admin_assigns_published_template(self):
        with app.app_context():
            _, emp, tpl = self._setup_admin_tpl_employee()
            emp_id, tpl_id = emp.id, tpl.id

        client = app.test_client()
        resp = client.post(
            "/assign",
            headers=as_user("admin@example.com"),
            data={
                "template_id": str(tpl_id),
                "user_id": str(emp_id),
                "start_date": "2026-07-06",
            },
        )
        assert resp.status_code == 302
        assert "/weeks/" in resp.headers["Location"]

        with app.app_context():
            emp = db.session.get(User, emp_id)
            assert emp.onboarding_plan_id is not None

            weeks = (
                Week.query.filter_by(onboarding_plan_id=emp.onboarding_plan_id)
                .order_by(Week.start_date)
                .all()
            )
            assert len(weeks) == 2  # one per template section
            assert weeks[0].start_date == date(2026, 7, 6)
            assert weeks[1].start_date == date(2026, 7, 13)  # +7 day offset
            for w in weeks:
                assert w.owner_user_id == emp_id
                assert Task.query.filter_by(week_id=w.id).count() == 2

    def test_assign_rejects_unpublished_template(self):
        with app.app_context():
            admin, emp, tpl = self._setup_admin_tpl_employee()
            tpl.status = TemplateStatusEnum.DRAFT.value
            db.session.commit()
            emp_id, tpl_id = emp.id, tpl.id

        client = app.test_client()
        resp = client.post(
            "/assign",
            headers=as_user("admin@example.com"),
            data={"template_id": str(tpl_id), "user_id": str(emp_id)},
        )
        assert resp.status_code == 400

    def test_assign_rejects_non_admin(self):
        with app.app_context():
            admin, emp, tpl = self._setup_admin_tpl_employee()
            make_user("mgr@example.com", "Max Manager", RoleEnum.MANAGER.value)
            db.session.commit()
            emp_id, tpl_id = emp.id, tpl.id

        client = app.test_client()
        resp = client.post(
            "/assign",
            headers=as_user("mgr@example.com"),
            data={"template_id": str(tpl_id), "user_id": str(emp_id)},
        )
        assert resp.status_code == 403


# ---------------------------------------------------------------------------
# 3. Task inline editing (add / status / notes / due date)
# ---------------------------------------------------------------------------
class TestTaskFlows:
    def _setup_owner_week(self):
        """A user who owns a week, plus an unrelated second user."""
        owner = make_user("owner@example.com", "Odie Owner", RoleEnum.USER.value)
        other = make_user("other@example.com", "Ola Other", RoleEnum.USER.value)
        plan = OnboardingPlan(name="Plan")
        db.session.add(plan)
        db.session.flush()
        owner.onboarding_plan_id = plan.id
        week = Week(
            title="Week 1",
            owner_user_id=owner.id,
            onboarding_plan_id=plan.id,
        )
        db.session.add(week)
        db.session.commit()
        return owner, other, week

    def _make_task(self, week_id, **kw):
        t = Task(week_id=week_id, title=kw.pop("title", "Task"), **kw)
        db.session.add(t)
        db.session.commit()
        return t

    def test_owner_adds_task(self):
        with app.app_context():
            owner, _, week = self._setup_owner_week()
            week_id = week.id

        client = app.test_client()
        resp = client.post(
            f"/weeks/{week_id}/tasks",
            headers=as_user("owner@example.com"),
            data={
                "goal": "Read handbook",
                "topic": "HR",
                "notes": "ch. 1-3",
                "due_date": "2026-07-10",
            },
        )
        assert resp.status_code == 302

        with app.app_context():
            t = Task.query.filter_by(week_id=week_id).one()
            assert t.title == "Read handbook"
            assert t.topic == "HR"
            assert t.notes == "ch. 1-3"
            assert t.due_date == date(2026, 7, 10)

    def test_non_owner_cannot_add_task(self):
        with app.app_context():
            _, other, week = self._setup_owner_week()
            week_id = week.id

        client = app.test_client()
        resp = client.post(
            f"/weeks/{week_id}/tasks",
            headers=as_user("other@example.com"),
            data={"goal": "Sneaky task"},
        )
        assert resp.status_code == 403

    def test_status_update_htmx_valid(self):
        with app.app_context():
            _, _, week = self._setup_owner_week()
            t = self._make_task(week.id)
            task_id = t.id

        client = app.test_client()
        resp = client.post(
            f"/tasks/{task_id}/status",
            headers={**as_user("owner@example.com"), "HX-Request": "true"},
            data={"status": StatusEnum.COMPLETE.value},
        )
        assert resp.status_code == 200  # returns display fragment

        with app.app_context():
            t = db.session.get(Task, task_id)
            assert t.status == StatusEnum.COMPLETE.value

    def test_status_update_htmx_invalid_value(self):
        with app.app_context():
            _, _, week = self._setup_owner_week()
            t = self._make_task(week.id)
            task_id = t.id

        client = app.test_client()
        resp = client.post(
            f"/tasks/{task_id}/status",
            headers={**as_user("owner@example.com"), "HX-Request": "true"},
            data={"status": "Done-ish"},
        )
        assert resp.status_code == 400

        with app.app_context():
            t = db.session.get(Task, task_id)
            assert t.status != "Done-ish"

    def test_notes_update_trims_whitespace(self):
        with app.app_context():
            _, _, week = self._setup_owner_week()
            t = self._make_task(week.id)
            task_id = t.id

        client = app.test_client()
        resp = client.post(
            f"/tasks/{task_id}/notes",
            headers={**as_user("owner@example.com"), "HX-Request": "true"},
            data={"notes": "  call IT about badge  "},
        )
        assert resp.status_code == 200

        with app.app_context():
            t = db.session.get(Task, task_id)
            assert t.notes == "call IT about badge"

    def test_due_date_update_and_clear(self):
        with app.app_context():
            _, _, week = self._setup_owner_week()
            t = self._make_task(week.id)
            task_id = t.id

        client = app.test_client()
        headers = {**as_user("owner@example.com"), "HX-Request": "true"}

        # set (US format accepted)
        resp = client.post(
            f"/tasks/{task_id}/due-date", headers=headers, data={"due_date": "07/10/26"}
        )
        assert resp.status_code == 200
        with app.app_context():
            assert db.session.get(Task, task_id).due_date == date(2026, 7, 10)

        # invalid -> 400, value unchanged
        resp = client.post(
            f"/tasks/{task_id}/due-date", headers=headers, data={"due_date": "sometime"}
        )
        assert resp.status_code == 400
        with app.app_context():
            assert db.session.get(Task, task_id).due_date == date(2026, 7, 10)

        # clear
        resp = client.post(
            f"/tasks/{task_id}/due-date", headers=headers, data={"clear": "1"}
        )
        assert resp.status_code == 200
        with app.app_context():
            assert db.session.get(Task, task_id).due_date is None

    def test_non_owner_cannot_edit_task(self):
        with app.app_context():
            _, other, week = self._setup_owner_week()
            t = self._make_task(week.id)
            task_id = t.id

        client = app.test_client()
        for method, path, data in [
            ("post", f"/tasks/{task_id}/status", {"status": "Complete"}),
            ("post", f"/tasks/{task_id}/notes", {"notes": "hax"}),
            ("post", f"/tasks/{task_id}/due-date", {"due_date": "2026-07-10"}),
        ]:
            resp = getattr(client, method)(
                path, headers=as_user("other@example.com"), data=data
            )
            assert resp.status_code == 403, path
